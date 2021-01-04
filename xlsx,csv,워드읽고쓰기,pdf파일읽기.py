import csv
import openpyxl as xl
import glob
import win32com.client

# xlsx파일 읽기
files=glob.glob('C:/Users/SNUH/Desktop/유일한/*.xlsx')
list=[]
for file in files:
    file=xl.load_workbook(file)
    file_a=file.active
    for r in file_a.rows:
        if r[0].value is None:
            continue
        list.append([])
        for c in r:
            list[-1].append(c.value)
        print(list[-1])

# csv파일읽기
files2=glob.glob('C:/Users/SNUH/Desktop/유일한/홍윤호교수님 연구/*.csv')
list=[]
for file in files2:
    file=open(file, 'r', encoding='utf-8')
    file_a=csv.reader(file)
    for line in file_a:
        print(line)
    file.close()

# csv파일 쓰기
for file in files2:
    file=open(file,'w',encoding='utf-8',newline='') #newline= 을 지정안하면 각 라인 뒤에 빈라인이 추가됨
    file_a=csv.writer(file)
    file_a.writerow([1,"aaa",False])
    file_a.writerow([2,"bbb",True])
file_a.close()

# word파일 읽기
from docx import Document
from docx.shared import Inches
import docx2txt

text=docx2txt.process("C:/Users/SNUH/Desktop/유일한/70873757 임재연 피타렉스 소견서.docx")
print(text)

# 워드 파일 쓰기
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells=1, cols=3)
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('C:/Users/SNUH/Desktop/유일한/demo.docx')

# PDF파일 읽기
from tika import parser
rawText=parser.from_file("C:/Users/SNUH/Desktop/유일한/논문/2018년 고혈압 진료지침.pdf")
rawList=rawText['content'].splitlines()
print(rawList)